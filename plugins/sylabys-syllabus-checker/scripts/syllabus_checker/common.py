from __future__ import annotations

import hashlib
import json
import os
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET


REPO_ROOT = Path(__file__).resolve().parents[2]
CONFIG_DIR = REPO_ROOT / "config"
DEFAULT_CONFIG_PATH = CONFIG_DIR / "syllabus-checker.json"
BUILD_DIR = REPO_ROOT / "build"
REPORTS_DIR = REPO_ROOT / "reports"

VERDICT_ORDER = {"FAIL": 0, "NEEDS_HUMAN": 1, "WARN": 2, "SKIP": 3, "PASS": 4}
SEVERITY_ORDER = {"CRITICAL": 0, "MAJOR": 1, "MINOR": 2, "WARN": 3}


@dataclass(frozen=True)
class AgentProfile:
    suite: str
    name: str
    role: str
    icon: str
    output: str


AGENTS: dict[str, AgentProfile] = {
    "STR": AgentProfile("STR", "syllabus-str", "Структурный контролёр", "structure", "str.json"),
    "FMT": AgentProfile("FMT", "syllabus-fmt", "Технический контролёр", "paintbrush", "fmt.json"),
    "OP": AgentProfile("OP", "syllabus-op", "Эксперт соответствия ОП", "graduation-cap", "op.json"),
    "RUP": AgentProfile("RUP", "syllabus-rup", "Эксперт учебного плана", "table-2", "rup.json"),
    "INT": AgentProfile("INT", "syllabus-int", "Аудитор согласованности", "list-checks", "int.json"),
    "TXT": AgentProfile("TXT", "syllabus-txt", "Редактор-вычитчик", "spell-check", "txt.json"),
}

DEFAULT_CONFIG: dict[str, Any] = {
    "fixtures": {
        "op": {
            "patterns": ["Приложение 6", "standarta op", "стандарта оп", "паспорт оп"],
            "extensions": [".docx"],
        },
        "rup": {
            "patterns": ["РУП", "учебный план", ".xlsx"],
            "extensions": [".xlsx"],
        },
    },
    "outputs": {
        "defaultRoot": "plugin",
        "buildDirName": "build",
        "reportsDirName": "reports",
    },
    "pdf": {
        "enabled": True,
        "envBrowser": "SYLABYS_PDF_BROWSER",
        "envSkip": "SYLABYS_SKIP_PDF",
        "browserNames": ["msedge", "chrome", "chromium", "google-chrome", "microsoft-edge"],
        "browserPaths": [
            "C:/Program Files/Microsoft/Edge/Application/msedge.exe",
            "C:/Program Files (x86)/Microsoft/Edge/Application/msedge.exe",
            "C:/Program Files/Google/Chrome/Application/chrome.exe",
            "C:/Program Files (x86)/Google/Chrome/Application/chrome.exe",
            "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
            "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
            "/usr/bin/google-chrome",
            "/usr/bin/chromium",
            "/usr/bin/chromium-browser",
            "/snap/bin/chromium",
        ],
    },
}


def load_config() -> dict[str, Any]:
    config = _merge_dicts({}, DEFAULT_CONFIG)
    for path in _config_paths():
        if not path.exists():
            continue
        try:
            config = _merge_dicts(config, json.loads(path.read_text(encoding="utf-8")))
        except Exception:
            continue
    return config


def _config_paths() -> list[Path]:
    paths = [DEFAULT_CONFIG_PATH, Path.cwd() / "sylabys-checker.json"]
    env_path = os.environ.get("SYLABYS_CHECKER_CONFIG")
    if env_path:
        paths.append(Path(env_path).expanduser())
    result: list[Path] = []
    seen: set[Path] = set()
    for path in paths:
        resolved = path.resolve()
        if resolved not in seen:
            result.append(resolved)
            seen.add(resolved)
    return result


def _merge_dicts(base: dict[str, Any], override: dict[str, Any]) -> dict[str, Any]:
    merged = dict(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = _merge_dicts(merged[key], value)
        else:
            merged[key] = value
    return merged


def config_list(config: dict[str, Any], *keys: str) -> list[str]:
    value: Any = config
    for key in keys:
        if not isinstance(value, dict):
            return []
        value = value.get(key)
    if isinstance(value, list):
        return [str(item) for item in value]
    if isinstance(value, str):
        return [value]
    return []


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def stable_id(value: str) -> str:
    return hashlib.sha1(value.encode("utf-8", "ignore")).hexdigest()[:12]


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(data, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )


def verdict(
    test_id: str,
    verdict_value: str,
    severity: str,
    *,
    discipline: str = "",
    location: str = "",
    expected: str = "",
    actual: str = "",
    evidence: str = "",
    recommendation: str = "",
    confidence: float = 1.0,
) -> dict[str, Any]:
    return {
        "testId": test_id,
        "verdict": verdict_value,
        "severity": severity,
        "discipline": discipline,
        "location": location,
        "expected": expected,
        "actual": actual,
        "evidence": evidence,
        "recommendation": recommendation,
        "confidence": confidence,
    }


def sorted_verdicts(items: Iterable[dict[str, Any]]) -> list[dict[str, Any]]:
    return sorted(
        items,
        key=lambda x: (
            x.get("testId", ""),
            VERDICT_ORDER.get(x.get("verdict", "PASS"), 9),
            x.get("discipline", ""),
            x.get("location", ""),
        ),
    )


def parse_first_number(text: str) -> float | None:
    match = re.search(r"(?<!\d)(\d+(?:[,.]\d+)?)(?!\d)", text or "")
    if not match:
        return None
    return float(match.group(1).replace(",", "."))


def _zip_xml_texts(path: Path, member_prefix: str) -> list[str]:
    out: list[str] = []
    with zipfile.ZipFile(path) as zf:
        for name in sorted(zf.namelist()):
            if not name.startswith(member_prefix) or not name.endswith(".xml"):
                continue
            root = ET.fromstring(zf.read(name))
            for node in root.iter():
                if node.tag.endswith("}t") and node.text:
                    out.append(node.text)
    return out


def read_docx(path: Path) -> dict[str, Any]:
    try:
        from docx import Document  # type: ignore

        doc = Document(str(path))
        paragraphs = [normalize_text(p.text) for p in doc.paragraphs if normalize_text(p.text)]
        tables: list[list[list[str]]] = []
        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                rows.append([normalize_text(cell.text) for cell in row.cells])
            tables.append(rows)
        text = "\n".join(paragraphs)
        return {"path": str(path), "paragraphs": paragraphs, "tables": tables, "text": text}
    except Exception as exc:
        try:
            fallback = _read_docx_zip(path)
            fallback["_fallback"] = f"python-docx unavailable: {type(exc).__name__}: {exc}"
            return fallback
        except Exception as fallback_exc:
            return {
                "path": str(path),
                "paragraphs": [],
                "tables": [],
                "text": "",
                "_extraction_error": f"{type(exc).__name__}: {exc}; fallback: {fallback_exc}",
            }


def _read_docx_zip(path: Path) -> dict[str, Any]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    paragraphs = []
    for paragraph in root.findall(".//w:p", ns):
        text = _word_text(paragraph, ns)
        if text:
            paragraphs.append(text)
    tables: list[list[list[str]]] = []
    for table in root.findall(".//w:tbl", ns):
        rows: list[list[str]] = []
        for row in table.findall("./w:tr", ns):
            cells = [_word_text(cell, ns) for cell in row.findall("./w:tc", ns)]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
    return {"path": str(path), "paragraphs": paragraphs, "tables": tables, "text": "\n".join(paragraphs)}


def _word_text(node: ET.Element, ns: dict[str, str]) -> str:
    return normalize_text("".join(text_node.text or "" for text_node in node.findall(".//w:t", ns)))


def inspect_docx_format(path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {
        "trackChanges": False,
        "comments": False,
        "highlightedRuns": 0,
        "fonts": {},
        "fontSizesHalfPt": {},
        "marginsCm": [],
    }
    try:
        with zipfile.ZipFile(path) as zf:
            names = set(zf.namelist())
            xml = zf.read("word/document.xml").decode("utf-8", "ignore")
            result["trackChanges"] = any(tag in xml for tag in ("<w:ins", "<w:del", "<w:moveFrom", "<w:moveTo"))
            result["comments"] = "word/comments.xml" in names or "<w:commentRangeStart" in xml
            result["highlightedRuns"] = xml.count("<w:highlight")
            for font in re.findall(r'w:ascii="([^"]+)"', xml):
                result["fonts"][font] = result["fonts"].get(font, 0) + 1
            for size in re.findall(r'<w:sz[^>]+w:val="([^"]+)"', xml):
                result["fontSizesHalfPt"][size] = result["fontSizesHalfPt"].get(size, 0) + 1
            for margin in re.findall(r"<w:pgMar[^>]+>", xml):
                vals = {}
                for key in ("top", "bottom", "left", "right"):
                    match = re.search(rf'w:{key}="(\d+)"', margin)
                    if match:
                        vals[key] = round(int(match.group(1)) / 567, 2)
                if vals:
                    result["marginsCm"].append(vals)
    except Exception as exc:
        result["_extraction_error"] = f"{type(exc).__name__}: {exc}"
    return result


def read_xlsx(path: Path) -> dict[str, Any]:
    try:
        from openpyxl import load_workbook  # type: ignore

        wb = load_workbook(str(path), data_only=True, read_only=True)
        sheets = {}
        for ws in wb.worksheets:
            rows = []
            for row in ws.iter_rows(values_only=True):
                values = [normalize_text(str(v)) if v is not None else "" for v in row]
                if any(values):
                    rows.append(values)
            sheets[ws.title] = rows
        return {"path": str(path), "sheets": sheets}
    except Exception as exc:
        fallback = _read_xlsx_zip(path)
        if fallback.get("sheets"):
            fallback["_fallback"] = f"openpyxl unavailable: {type(exc).__name__}: {exc}"
            return fallback
        return {"path": str(path), "sheets": {}, "_extraction_error": f"{type(exc).__name__}: {exc}"}


def _read_xlsx_zip(path: Path) -> dict[str, Any]:
    ns = {
        "a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    }
    try:
        with zipfile.ZipFile(path) as zf:
            shared = _xlsx_shared_strings(zf, ns)
            workbook = ET.fromstring(zf.read("xl/workbook.xml"))
            rels = ET.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
            rel_targets = {
                rel.attrib.get("Id"): rel.attrib.get("Target", "")
                for rel in rels.findall("rel:Relationship", ns)
            }
            sheets: dict[str, list[list[str]]] = {}
            for sheet in workbook.findall(".//a:sheet", ns):
                title = sheet.attrib.get("name") or "Sheet"
                rid = sheet.attrib.get(f"{{{ns['r']}}}id")
                target = rel_targets.get(rid, "")
                if not target:
                    continue
                member = "xl/" + target.lstrip("/")
                if member not in zf.namelist():
                    member = "xl/worksheets/" + Path(target).name
                rows = _xlsx_sheet_rows(zf.read(member), shared, ns)
                if rows:
                    sheets[title] = rows
            return {"path": str(path), "sheets": sheets}
    except Exception as exc:
        return {"path": str(path), "sheets": {}, "_extraction_error": f"{type(exc).__name__}: {exc}"}


def _xlsx_shared_strings(zf: zipfile.ZipFile, ns: dict[str, str]) -> list[str]:
    if "xl/sharedStrings.xml" not in zf.namelist():
        return []
    root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
    values = []
    for item in root.findall("a:si", ns):
        values.append(normalize_text("".join(node.text or "" for node in item.findall(".//a:t", ns))))
    return values


def _xlsx_sheet_rows(xml: bytes, shared: list[str], ns: dict[str, str]) -> list[list[str]]:
    root = ET.fromstring(xml)
    rows = []
    for row in root.findall(".//a:row", ns):
        values = []
        for cell in row.findall("a:c", ns):
            value_node = cell.find("a:v", ns)
            inline_nodes = cell.findall(".//a:t", ns)
            value = ""
            if cell.attrib.get("t") == "s" and value_node is not None:
                idx = int(value_node.text or 0)
                value = shared[idx] if 0 <= idx < len(shared) else ""
            elif inline_nodes:
                value = "".join(node.text or "" for node in inline_nodes)
            elif value_node is not None:
                value = value_node.text or ""
            values.append(normalize_text(value))
        if any(values):
            rows.append(values)
    return rows


def find_fixture(
    patterns: list[str],
    root: Path = REPO_ROOT,
    *,
    roots: Iterable[Path] | None = None,
    extensions: Iterable[str] | None = None,
    recursive: bool = False,
) -> Path | None:
    search_roots = list(roots or [root])
    ext_set = {ext.casefold() for ext in (extensions or []) if ext}
    candidates: list[Path] = []
    seen: set[Path] = set()
    for search_root in search_roots:
        if not search_root.exists():
            continue
        iterator = search_root.rglob("*") if recursive else search_root.iterdir()
        for path in iterator:
            if not path.is_file():
                continue
            resolved = path.resolve()
            if resolved in seen:
                continue
            if ext_set and path.suffix.casefold() not in ext_set:
                continue
            candidates.append(path)
            seen.add(resolved)
    for pattern in patterns:
        needle = pattern.casefold()
        for path in candidates:
            if needle in path.name.casefold():
                return path
    return None
